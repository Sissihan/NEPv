# -*- coding: utf-8 -*-
"""Generate NEPv Lens installation, deployment, and operation manual (.docx) in English."""

from __future__ import annotations

import json
import subprocess
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "NEPv-Lens-Installation-Deployment-and-Operation-Manual.docx"
OUT_ALT = ROOT / "docs" / "NEPv-Lens-Manual-With-Screenshots.docx"
SHOT_DIR = ROOT / "docs" / "manual-screenshots"
MANIFEST = SHOT_DIR / "manifest.json"
CAPTURE_SCRIPT = ROOT / "scripts" / "capture_manual_screenshots.py"


def item_label(item: dict) -> str:
    return item.get("label") or item.get("label_zh", "")


def item_section(item: dict) -> str:
    return item.get("section") or item.get("section_zh", "")


def item_action(item: dict) -> str:
    return item.get("action") or item.get("action_zh", "")


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def p(doc: Document, text: str, bold: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    para.paragraph_format.space_after = Pt(6)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def code(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    para.paragraph_format.left_indent = Cm(0.5)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
    for r_idx, row in enumerate(rows):
        cells = t.rows[r_idx + 1].cells
        for c_idx, text in enumerate(row):
            cells[c_idx].text = text
    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.is_file():
        p(doc, f"[Screenshot missing: {image_path.name}. Run scripts/capture_manual_screenshots.py first.]")
        return
    doc.add_picture(str(image_path), width=Inches(5.8))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    doc.add_paragraph()


def ensure_screenshots() -> list[dict]:
    if MANIFEST.is_file():
        return json.loads(MANIFEST.read_text(encoding="utf-8"))
    print("Screenshots not found — running capture script…")
    subprocess.run([sys.executable, str(CAPTURE_SCRIPT)], check=True, cwd=ROOT)
    if not MANIFEST.is_file():
        raise FileNotFoundError(f"Capture failed: {MANIFEST}")
    return json.loads(MANIFEST.read_text(encoding="utf-8"))


def build_part1_install(doc: Document) -> None:
    h(doc, "Part I — Installation & Deployment", 1)

    h(doc, "1. Purpose", 2)
    p(
        doc,
        "This manual is for administrators, reviewers, and learners. It explains how to install, "
        "deploy, verify, and troubleshoot NEPv Lens on Windows, macOS, or Linux. The product is a "
        "static front-end web app (Vite 8 + React 18); there is no backend database.",
    )

    h(doc, "2. System requirements (pinned versions)", 2)
    table(
        doc,
        ["Item", "Version / requirement"],
        [
            ["Operating system", "Windows 10/11, macOS 12+, mainstream Linux"],
            ["Node.js", "≥ 18.17.0 (.nvmrc / engines)"],
            ["React", "18.2.0 (package.json)"],
            ["Vite", "8.x (@vitejs/plugin-react 6.x)"],
            ["Package manager", "npm (bundled with Node)"],
            ["Browser", "Chrome / Edge / Firefox (last two major versions)"],
            ["Optional: Docker", "Docker Desktop or Engine + Compose"],
            ["Manual screenshot tooling", "Python 3 + python-docx + playwright (chromium)"],
            ["Disk space", "≥ 500 MB (including node_modules)"],
        ],
    )

    h(doc, "3. Obtain source code", 2)
    p(doc, "Unpack or clone the project. The root must contain package.json, src/, docs/, and Dockerfile.")

    h(doc, "4. Local Node.js workflow", 2)
    code(doc, "npm install")
    code(doc, "npm run dev          # http://localhost:5173")
    code(doc, "npm run build && npm run preview")

    h(doc, "5. Docker deployment", 2)
    code(doc, "docker-compose up --build")

    h(doc, "6. Acceptance checklist", 2)
    table(
        doc,
        ["Check", "Command / action", "Pass criterion"],
        [
            ["Dependencies", "npm install", "No errors"],
            ["Unit tests", "npm test", "Residual baseline tests pass (≤ 1e-6 error)"],
            ["Type check", "npx tsc --noEmit", "No errors"],
            ["Production build", "npm run build", "dist/ created"],
            ["Bundle size", "inspect dist/assets", "Main JS gzip ≲ 150 KB"],
            ["UI smoke test", "open localhost:5173", "Intro band + lab interactive"],
            ["Freeze pitfall", "check “Pitfall: freeze A”", "⚠ badge + frozen spectrum column"],
            ["AI tutor", "ask a preset question", "Rule-based answer from live state"],
            ["Manual screenshots", "python scripts/capture_manual_screenshots.py", "docs/manual-screenshots/ has 45+ PNG"],
        ],
    )

    h(doc, "7. Regenerate this manual", 2)
    p(doc, "From the project root:")
    code(doc, "npm run build")
    code(doc, "python scripts/capture_manual_screenshots.py")
    code(doc, "python scripts/generate_manual_docx.py")


def build_part2_operation(doc: Document, manifest: list[dict]) -> None:
    h(doc, "Part II — Operation manual (with screenshots)", 1)

    h(doc, "1. Product overview", 2)
    p(
        doc,
        "NEPv Lens is a teaching demo for eigenvector-dependent nonlinear eigenvalue problems: "
        "A(x)·x = λ·x with A depending on x. It does not claim industrial global NEPv solving.",
    )

    h(doc, "2. Page map", 2)
    bullet(doc, "Definition & Compare — side-by-side intro band (read-only).")
    bullet(doc, "Lab — compass, heatmap/spectrum, polar residual plot, iteration lab, setup drawer.")
    bullet(doc, "AI Math Tutor — full-width section under the lab (status, presets, Q&A).")
    bullet(doc, "Pitfalls — freeze-A pitfall card and FAQ grid.")
    bullet(doc, "References — collapsible bibliography.")
    bullet(doc, "Locale — English (default for this manual), Simplified Chinese, Traditional Chinese.")

    h(doc, "3. Control inventory", 2)
    p(
        doc,
        f"The table lists {len(manifest)} UI regions and controls captured for version 0.1.x. "
        "Screenshots follow the teaching flow top-to-bottom.",
    )
    table(
        doc,
        ["#", "Control / region", "Type", "Section", "Summary"],
        [
            [
                str(i + 1),
                item_label(item),
                item["control_type"],
                item_section(item),
                item_action(item)[:90] + ("…" if len(item_action(item)) > 90 else ""),
            ]
            for i, item in enumerate(manifest)
        ],
    )

    h(doc, "4. Step-by-step (one screenshot per control)", 2)
    p(doc, "Each subsection matches a button, input, or display region with a live UI capture (English UI).")

    current_section = ""
    for i, item in enumerate(manifest, start=1):
        sec = item_section(item)
        if sec != current_section:
            current_section = sec
            h(doc, sec, 3)

        h(doc, f"{i}. {item_label(item)}", 4)
        table(
            doc,
            ["Property", "Details"],
            [
                ["ID / selector", item.get("control_id", "") + (f" ({item['selector']})" if item.get("selector") else "")],
                ["Type", item["control_type"]],
                ["How to use", item_action(item)],
            ],
        )
        img = SHOT_DIR / item["file"]
        add_figure(doc, img, f"Figure {i}: {item_label(item)}")

    h(doc, "5. Recommended learning path", 2)
    bullet(doc, "Read Definition & Compare (~3 min) to contrast linear EVP and NEPv.")
    bullet(doc, "In the lab, pick Model A; drag the compass and α; watch heatmap, polar plot, and spectrum (~5 min).")
    bullet(doc, "Enable “Pitfall: freeze A” and compare frozen vs true A(x) (~2 min).")
    bullet(doc, "Switch to Model B; tune β₁–β₃ and x₃; use Play / Step / Reset in the iteration lab (~5 min).")
    bullet(doc, "Open AI Tutor presets and ask why iteration stalls (~2 min).")
    bullet(doc, "Expand References; switch locale to verify i18n (~2 min).")

    h(doc, "6. Mathematical notes & interaction", 2)
    bullet(doc, "Relative residual banner: r_rel(x,λ) = ||A(x)x − λx||₂ / (||A(x)x||₂ + ||λx||₂).")
    bullet(doc, "Toy models include verified baselines (error ≤ 1e-6 in unit tests).")
    bullet(doc, "Compass enforces ||x|| = 1; singular x = 0 disables controls.")
    bullet(doc, "Complex eigenvalues appear as conjugate pairs in the spectrum bars.")
    bullet(doc, "Freeze mode is labeled as a pitfall (linear EVP at fixed x₀, not true NEPv).")
    bullet(doc, "Auto-iteration stops after 200 steps if r > 1e-3; use Reset to reference.")
    bullet(doc, "Mobile (≤ 375px): numeric x₁/x₂ fields and stacked observe layout.")

    h(doc, "7. Limitations & disclaimer", 2)
    bullet(doc, "Small toy models only (n ≤ 4); not a production NEPv solver.")
    bullet(doc, "Iteration may fail to converge — intentional teaching point.")
    p(
        doc,
        "Spectra show instantaneous linear eigenvalues of the current A(x), not a global NEPv spectrum. "
        "AI responses are rule-based templates (AI-assisted, human-verified), not a certified LLM backend.",
    )

    h(doc, "8. AI tutor verification checklist", 2)
    table(
        doc,
        ["Case", "Expected behavior", "Status"],
        [
            ["High residual, many steps", "Suggest lower α, change x, turn off freeze", "Verified"],
            ["Freeze on", "State clearly: solving a linear EVP at x₀", "Verified"],
            ["Large α and r > 0.05", "Suggest α ≈ 0.4–0.6", "Verified"],
            ["x near zero", "Warn that x cannot be zero", "Verified"],
        ],
    )


def build_document() -> Document:
    manifest = ensure_screenshots()
    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("NEPv Lens\nInstallation, Deployment & Operation Manual")
    r.bold = True
    r.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"Version: 0.1.x (UI refresh — intro band, lab shell, centered AI tutor)\n"
        f"Date: {date.today().isoformat()}\n"
        f"Screenshots: {len(manifest)} (English UI)\n"
        "Teaching demo: nonlinear eigenvalue problems with eigenvector dependency\n"
        "AI-assisted, human-verified"
    )

    doc.add_page_break()
    build_part1_install(doc)
    doc.add_page_break()
    build_part2_operation(doc, manifest)

    doc.add_paragraph()
    p(doc, "— End of document —", bold=True)
    return doc


def save_doc(doc: Document, path: Path) -> Path | None:
    try:
        doc.save(path)
        return path
    except PermissionError:
        print(f"Warning: cannot write {path} (file may be open in Word). Skipped.")
        return None


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    written: list[Path] = []
    for path in (OUT, OUT_ALT):
        saved = save_doc(doc, path)
        if saved:
            written.append(saved)
    if not written:
        raise PermissionError("All target .docx paths are locked — close Word and retry.")
    for path in written:
        print(f"Generated: {path}")


if __name__ == "__main__":
    main()
